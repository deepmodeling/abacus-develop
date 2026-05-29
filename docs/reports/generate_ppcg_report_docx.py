#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Generate a simple editable .docx from the PPCG Markdown report.

Design goal:
- Keep formatting clean and editable (Headings + paragraphs + bullet lists).
- Minimal markdown parsing (headings, blockquotes, unordered lists, code fences).

Usage:
  python3 docs/reports/generate_ppcg_report_docx.py \
      docs/reports/PPCG_算法实现报告.md \
      docs/reports/PPCG_算法实现报告.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)\s*$")
LIST_RE = re.compile(r"^\s*[-*]\s+(.*)\s*$")
INLINE_MATH_RE = re.compile(r"\$(.+?)\$")


def latex_to_unicode(expr: str) -> str:
    # Minimal, pragmatic conversion for this report.
    # Goal: readable equations in Word without requiring a full LaTeX->OMML converter.
    s = expr

    # Common LaTeX commands used in the report
    replacements = {
        r"\\lambda": "λ",
        r"\\Lambda": "Λ",
        r"\\dagger": "†",
        r"\\times": "×",
        r"\\approx": "≈",
        r"\\leftarrow": "←",
        r"\\in": "∈",
        r"\\mathbb{C}": "ℂ",
        r"\\mathbb{R}": "ℝ",
        r"\\mathbb{Z}": "ℤ",
        r"\\mathbb{N}": "ℕ",
    }
    for k, v in replacements.items():
        s = s.replace(k, v)

    # Handle ^\dagger / ^{\dagger}
    s = s.replace(r"^\\dagger", "†")
    s = s.replace(r"^{\\dagger}", "†")

    # Superscripts for simple integer exponents: ^{-1}, ^{2}, ^2
    sup_map = str.maketrans({
        "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
        "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
        "+": "⁺", "-": "⁻",
    })

    def supify(num: str) -> str:
        return "".join(ch.translate(sup_map) for ch in num)

    s = re.sub(r"\^\{([+-]?\d+)\}", lambda m: supify(m.group(1)), s)
    s = re.sub(r"\^([+-]?\d)", lambda m: supify(m.group(1)), s)

    # Remove LaTeX spacing commands we don't need
    s = s.replace(r"\\,", " ")
    s = s.replace(r"\\;", " ")

    # Strip outer braces in simple cases
    s = s.replace("{", "").replace("}", "")

    return s


def append_omml_inline(paragraph, expr: str) -> None:
    """Append an OMML inline equation to an existing paragraph."""
    omath = OxmlElement("m:oMath")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    # Preserve spaces inside equation text
    t.set(qn("xml:space"), "preserve")
    t.text = latex_to_unicode(expr)
    r.append(t)
    omath.append(r)
    paragraph._p.append(omath)


def add_math_paragraph(doc: Document, expr: str) -> None:
    """Add a standalone display equation paragraph (OMML)."""
    p = doc.add_paragraph("")
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.set(qn("xml:space"), "preserve")
    t.text = latex_to_unicode(expr)
    r.append(t)
    omath.append(r)
    omath_para.append(omath)
    p._p.append(omath_para)


def add_paragraph_with_inline_math(doc: Document, text: str, style: str | None = None):
    """Create a paragraph and render any $...$ as OMML equations."""
    p = doc.add_paragraph("", style=style) if style else doc.add_paragraph("")
    idx = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > idx:
            p.add_run(text[idx:m.start()])
        append_omml_inline(p, m.group(1))
        idx = m.end()
    if idx < len(text):
        p.add_run(text[idx:])
    return p


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    in_code = False
    code_lines: list[str] = []

    in_display_math = False
    display_math_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip("\n")

        # Display math blocks with $$ ... $$ (single or multi-line)
        if not in_code and line.strip().startswith("$$"):
            if not in_display_math:
                in_display_math = True
                display_math_lines = []
                # Handle single-line $$expr$$
                if line.strip().endswith("$$") and len(line.strip()) > 4:
                    expr = line.strip()[2:-2].strip()
                    add_math_paragraph(doc, expr)
                    in_display_math = False
                    display_math_lines = []
                continue
            else:
                # End of multi-line display math
                in_display_math = False
                expr = "\n".join(display_math_lines).strip()
                add_math_paragraph(doc, expr)
                display_math_lines = []
                continue

        if in_display_math:
            # Strip a trailing $$ on the last line if user wrote it that way
            if line.strip().endswith("$$"):
                display_math_lines.append(line.strip()[:-2].rstrip())
                in_display_math = False
                expr = "\n".join(display_math_lines).strip()
                add_math_paragraph(doc, expr)
                display_math_lines = []
            else:
                display_math_lines.append(line)
            continue

        # Code fences
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Empty line -> spacing
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Blockquote -> normal paragraph
        if line.lstrip().startswith(">"):
            content = line.lstrip()[1:].lstrip()
            add_paragraph_with_inline_math(doc, content)
            continue

        # Headings
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            # Word heading levels: 0=Title, 1..9 are Heading 1..9
            if level == 1:
                doc.add_heading(title, level=0)
            else:
                doc.add_heading(title, level=min(level - 1, 9))
            continue

        # Unordered list
        m = LIST_RE.match(line)
        if m:
            add_paragraph_with_inline_math(doc, m.group(1).strip(), style="List Bullet")
            continue

        # Default paragraph
        add_paragraph_with_inline_math(doc, line)

    # If file ended inside a code block, flush it.
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: generate_ppcg_report_docx.py <input.md> <output.docx>")
        return 2

    md_path = Path(argv[1])
    docx_path = Path(argv[2])

    if not md_path.exists():
        print(f"Input markdown not found: {md_path}")
        return 1

    convert(md_path, docx_path)
    print(f"Wrote: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
